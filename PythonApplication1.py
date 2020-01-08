import os

import docx
from lxml import html
import xlrd
import zipfile
import win32com.client


def copy_tables(doc_path, count):
    '''
    有几个问题，就把word模板中的空白表格复制几遍
    '''
    app = win32com.client.DispatchEx('Word.Application')
    app.Visable = 0
    doc = app.Documents.Open(doc_path)
    doc.Range(0, 0).Select()
    wordSel = app.Selection
    wordSel.WholeStory()
    wordSel.Copy()
    for i in range(count):
        wordSel.PasteAndFormat(16)
    doc.Close()


def isfile_exist(file_path):
    '''
    判断目标文件是否存在
    '''
    if not os.path.isfile(file_path):
        print("It's not a file or no such file exist! %s" % file_path)
        return False
    else:
        return True


def change_file_name(file_path, new_type='.zip'):
    '''
    将文件后缀名强制修改为zip
    '''
    if not isfile_exist(file_path):
        return ''
    extend = os.path.basename(file_path)[-5:]
    if extend != '.xlsx':
        print("It's not a xlsx file! %s" % file_path)
        return False
    file_name = os.path.basename(file_path)
    new_name = str(file_name.split('.')[0]) + new_type
    dir_path = os.path.dirname(file_path)
    new_path = os.path.join(dir_path, new_name)

    if os.path.exists(new_path):
        os.remove(new_path)

    os.rename(file_path, new_path)
    return new_path


def unzip_file(zipfile_path):
    '''
    解压指定zip文件到同名文件夹下
    '''
    if not isfile_exist(zipfile_path):
        return False
    if os.path.splitext(zipfile_path)[1] != '.zip':
        print("It's not a zip file! %s" % zipfile_path)

    file_zip = zipfile.ZipFile(zipfile_path, 'r')
    file_name = os.path.basename(zipfile_path)
    zipdir = os.path.join(os.path.dirname(zipfile_path), str(file_name.split('.')[0]))
    for files in file_zip.namelist():
        file_zip.extract(files, os.path.join(zipfile_path, zipdir))
    file_zip.close()
    return True


def fill_pic(excel_file_path, word_file_path):
    '''
    将xlsx文件中的图片添加到指定docx文件的表格中
    '''
    file = docx.Document(word_file_path)
    init_num = 0

    for i in range(len(file_xlsx.sheets())):
        xml = html.etree.parse(
            excel_file_path[:-5] + '/xl/drawing' + str(i + 1) + '.xml', html.etree.HTMLParser())
        # result_col = xml.xpath('//from//col//text()')
        result_row = xml.xpath('//from//row//text()')
        result_name = xml.xpath('//cnvpr/@name')
        result_id = xml.xpath('//cnvpr/@id')
        result_embed = xml.path('//blip/@*[local-name()="r:embed"]')

        xml2 = html.etree.parse(
            excel_file_path[:-5] + 'xl/drawings/_rels/drawing' + str(i + 1) + '.xml.rels', html.etree.HTMLParser())
        result_target=xml2.xpath('//*/@target')
        result_id2=xml2.xpath('//*/@id')

        pic=[]
        for i in range(len(result_embed)):
            for j in range(len(result_id2)):
                if result_embed[i] == result_id2[j]:
                    pic.append(result_target[j])
                    file.tables[2 * ((int(result_row[j]) + init_num)-1)].rows[8].cells[i].paragraphs[0].add_run().add_picture(excel_file_path[:-5]+'xl'+pic)

        init_num += len(result_embed)

    file.save(word_file_path)
    return ''


def fill_table(file_docx, file_xlsx, word_file_path):
    '''
    将xlsx文件中各行内容填入docx文件对应表格单元格内
    '''
    num=0
    for i in range(len(file_xlsx.sheets())):
        for j in range(file_xlsx.sheet()[i].nrows - 1):
            file_docx.tables[2 * num].rows[2].cells[1].text=file_xlsx.sheet()[i].row_values(j+1)[1]

            num += 1
        file_docx.save(word_file_path)


def table_count(file_xlsx):
    '''
    计算不同sheet页下的问题总数
    '''
    sheet=0
    count=0
    for datas in file_xlsx.sheets():
        sheet += 1
        count += datas.nrows - 1
    return count


if __name__ == '__main__':
    excel_file_path=''
    word_file_path=''

    file_xlsx=xlrd.open_workbook(excel_file_path)

    table_count=table_count(file_xlsx)

    copy_tables(word_file_path, table_count)

    file_docx=docx.Document(word_file_path)

    fill_table(file_docx, file_xlsx, word_file_path)

    zip_file_path=change_file_name(excel_file_path)
    if zip_file_path != '':
        if unzip_file(zip_file_path):
            fill_pic(excel_file_path, word_file_path)
